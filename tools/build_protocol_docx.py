"""Generate the Control4 UDP protocol Word manual."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "docs" / "Control4-Audio-UDP-Protocol.docx"


def shade_header(row) -> None:
    for cell in row.cells:
        tc = cell._tePr if False else cell._tc
        tc_pr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F4E79")
        shd.set(qn("w:val"), "clear")
        tc_pr.append(shd)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True


def set_cell_font(cell, size=10, mono=False) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.font.name = "Consolas" if mono else "Calibri"


def add_table(doc, headers, rows, mono_cols=None) -> None:
    mono_cols = set(mono_cols or [])
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        hdr.cells[i].text = text
        set_cell_font(hdr.cells[i], 10)
    shade_header(hdr)
    for r_i, row in enumerate(rows):
        for c_i, value in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = str(value)
            set_cell_font(cell, 10, mono=c_i in mono_cols)
    doc.add_paragraph()


def add_code(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(31, 78, 121)


def heading_style(doc) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for name, size in (("Title", 28), ("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)):
        s = doc.styles[name]
        s.font.color.rgb = RGBColor(31, 78, 121)
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = True


def note(doc, text: str, title="Note") -> None:
    p = doc.add_paragraph()
    r = p.add_run(f"{title}. ")
    r.bold = True
    r.font.color.rgb = RGBColor(192, 80, 77)
    p.add_run(text)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    heading_style(doc)

    title = doc.add_paragraph("Control4 Ethernet Audio", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub = doc.add_paragraph("UDP command and response reference")
    sub.runs[0].font.size = Pt(16)
    sub.runs[0].font.color.rgb = RGBColor(89, 89, 89)

    meta = doc.add_paragraph()
    meta.add_run("Unofficial reverse-engineered protocol\n").italic = True
    meta.add_run(
        "For C4-AMP108-1B, C4-16AMP3-B, and AVM-16S1-B / C4-16ZAMSV3-B class audio switch. "
        "Control4 has not published this API. Commands were recovered from Director packet logs "
        "and live probes on UDP port 8750."
    )

    doc.add_heading("1. Purpose and scope", 1)
    doc.add_paragraph(
        "This manual describes how to talk directly to Control4 matrix amplifiers and the "
        "16×16 stereo audio matrix switch over the LAN, without a Control4 driver running in "
        "the client. Director (EA-5 and similar) uses the same language. Home Assistant, a "
        "UDP probe, or a small script can send the same packets."
    )
    doc.add_paragraph(
        "A Triad audio matrix (TCP, binary) is a different product and is not covered here."
    )

    doc.add_heading("1.1 Hardware used to confirm this document", 2)
    add_table(
        doc,
        ["Role", "Composer type", "Example IP", "Firmware / identity"],
        [
            [
                "Director",
                "c4:control4_ea5",
                "192.168.68.117",
                "Subscribes to ethernet events on the switch",
            ],
            [
                "8-zone amplifier",
                "c4:v3_16chanamp:c4-16amp3-b",
                "192.168.68.93",
                '03.26.52  /  c4.sy.info "v3_16chanamp:c4-16amp3-b:…"',
            ],
            [
                "16×16 audio switch",
                "c4:v3_avswitch:avm-16s1-b",
                "192.168.68.80",
                '03.26.53  /  c4.sy.info "c4:v3_avswitch:avm-16s1-b:…"',
            ],
            [
                "4-zone amplifier (earlier capture)",
                "C4-AMP108-1B",
                "192.168.2.169 (historical)",
                '03.24.45 — source of igain / eq / psave traces',
            ],
        ],
        mono_cols={2, 3},
    )
    note(
        doc,
        "IPs change with DHCP. Always identify the chassis with c4.sy.fwv and c4.sy.info before sending SET commands.",
        "Warning",
    )

    doc.add_heading("2. Transport", 1)
    add_table(
        doc,
        ["Item", "Value"],
        [
            ["Protocol", "UDP"],
            ["Port", "8750"],
            ["Encoding", "US-ASCII"],
            ["Line ending", "CR LF  (bytes 0D 0A)"],
            ["Who can talk", "The device replies to the UDP source address of each packet. Director does not have to be the only client."],
            ["Discovery", "SDDP exists for Composer identify; it is not required for these commands."],
        ],
    )

    doc.add_heading("3. Packet framing", 1)
    doc.add_paragraph("Every packet is one line:")
    add_code(doc, "{type}{sequence} {body}\\r\\n")
    doc.add_paragraph(
        "Type is two characters. Sequence is four hex digits (0000–FFFF). The device copies "
        "the sequence into the matching 0r reply so the client can pair requests with answers."
    )
    add_table(
        doc,
        ["Type", "Direction", "Meaning", "Example"],
        [
            ["0s", "Client → device", "SET / do something", "0se76c c4.amp.out 02 01"],
            ["0g", "Client → device", "GET / query", "0g630b c4.sy.fwv"],
            ["0r", "Device → client", "Reply to that sequence", '0r630b 000 c4.sy.fwv "03.26.52"'],
            ["0t", "Device → subscribers", "Unsolicited status", "0t0040 sa c4.amp.mute 02 00"],
        ],
        mono_cols={0, 3},
    )
    doc.add_paragraph(
        "Unsolicited 0t packets usually contain the token sa (“status available”) after the "
        "sequence, then the same command body a SET would have used. Director uses c4.sy.sub "
        "\"ethernet\" so it receives those events. A one-shot probe still gets 0r replies "
        "without subscribing."
    )

    doc.add_heading("3.1 Reply status tokens", 2)
    add_table(
        doc,
        ["Token", "Meaning"],
        [
            ["000", "Success. GET replies usually continue with the command name and data."],
            ["v01", "Accepted / value OK. Common on some SET replies. Also seen when a SET is ignored (for example volume above 100 on the switch)."],
            ["n01", "Not supported on this model or firmware."],
            ["e00", "Error — typically bad arguments (mute without 00/01, GET c4.asw.out with no zone)."],
            ["(empty)", "The device heard the packet and returned 0r{seq} with no token. Typical for a wrong command family (for example c4.amp.* sent to the switch)."],
        ],
        mono_cols={0},
    )

    doc.add_heading("3.2 Numbers in command bodies", 2)
    doc.add_paragraph(
        "Zone, input, mute, and most levels are two-digit hexadecimal bytes, 1-based for jacks "
        "(01 = jack 1). Output 16 is hex 10, not decimal 16. Decimal 16 as a zone argument is a "
        "different slot and can return v01 without meaning “output 16”."
    )
    add_code(doc, "Zone 1  →  01\nZone 8  →  08\nZone 16 →  10\nDisconnect / off →  00")

    doc.add_heading("4. System commands (both products)", 1)
    doc.add_paragraph("These use the c4.sy.* namespace and work on both the amplifier and the audio switch.")
    add_table(
        doc,
        ["Get/Set", "Body", "Typical reply"],
        [
            ["GET", "c4.sy.fwv", '000 c4.sy.fwv "03.26.52"'],
            ["GET", "c4.sy.info", '000 c4.sy.info "v3_16chanamp:c4-16amp3-b:…"'],
            ["GET", "c4.sy.sub", '000 c4.sy.sub "ethernet 192.168.68.117:-21887"'],
            ["SET", 'c4.sy.sub "ethernet"', "000 (Director subscribe)"],
            ["GET", "c4.sy.mac / c4.sy.ver / c4.sy.ip", "n01 on the AVM-16S1-B"],
            ["GET", "c4.sy.afwv", "n01 on AMP108"],
        ],
        mono_cols={0, 1, 2},
    )

    doc.add_heading("5. Matrix amplifier — namespace c4.amp", 1)
    doc.add_paragraph(
        "Applies to C4-AMP108-1B (4 stereo zones) and C4-16AMP3-B (8 stereo zones). "
        "GET lists return one hex byte per zone, left to right = zone 1 … N."
    )

    doc.add_heading("5.1 Routing, mute, and volume", 2)
    add_table(
        doc,
        ["Get/Set", "Body", "Notes"],
        [
            ["SET", "c4.amp.out {zone} {input}", "Route an analog/digital input to a speaker zone. input 00 disconnects the zone (off)."],
            ["GET", "c4.amp.ain", "Current input per zone. 00 = disconnected. AMP108: 4 values. 16AMP3: 8 values."],
            ["SET", "c4.amp.mute {zone} 00|01", "01 muted, 00 unmuted. Second byte is required (e00 without it)."],
            ["GET", "c4.amp.amut", "Mute flags per zone."],
            ["SET", "c4.amp.chvol {zone} {hex}", "Set playback volume. See volume scale below."],
            ["GET", "c4.amp.avol", "Volume hex per zone."],
            ["GET", "c4.amp.digi {input}", "Director queries this before routing; digital-input related."],
            ["GET", "c4.amp.vlim", "Volume limits per zone. 00 = none."],
            ["Status", "c4.amp.chlim {zone} 00", "Unsolicited after volume ramps."],
        ],
        mono_cols={0, 1},
    )
    note(
        doc,
        "Do not send c4.amp.chvolmax from a volume slider. Public Home Assistant implementations document a firmware bug: that command can snap live volume to the cap.",
        "Warning",
    )

    doc.add_heading("5.2 Amplifier volume scale", 2)
    doc.add_paragraph(
        "Amplifier volume is not 0–100 stored as hex. It is percent plus 155, as one byte:"
    )
    add_code(doc, "hex_byte = percent + 155\n  0%  →  9b  (155)\n 12%  →  a7  (167)\n100% →  ff  (255)  if the firmware allows it")
    doc.add_paragraph(
        "This matches AMP108 Director captures (chvol … a7 while lowering volume through a2, a0, 9f) "
        "and independent Home Assistant matrix-amp code. Older community scripts used +160 and read about 5% high."
    )
    doc.add_paragraph(
        "Live 16AMP3 example (all zones disconnected, leftover volume registers):"
    )
    add_code(doc, "GET c4.amp.avol  →  b0 a4 ba a4 a4 ad a4 9c\n≈ 21%, 9%, 31%, 9%, 9%, 18%, 9%, 1%")

    doc.add_heading("5.3 Tone, EQ, and source leveling", 2)
    add_table(
        doc,
        ["Get/Set", "Body", "Composer equivalent"],
        [
            ["GET", "c4.amp.abss / atrb / abal", "Bass, treble, balance per zone (signed hex)."],
            ["SET", "c4.amp.bassgain {zone} {signed}", "Shelving bass gain."],
            ["SET", "c4.amp.trebgain {zone} {signed}", "Shelving treble gain."],
            ["SET", "c4.amp.bassfreq {zone} {hex}", "Bass shelf frequency (Composer e.g. 125 Hz)."],
            ["SET", "c4.amp.trebfreq {zone} {hex}", "Treble shelf frequency (Composer e.g. 5 kHz)."],
            ["Status", "c4.amp.tone {zone} g f g f", "Unsolicited after tone SET: bass gain, bass freq, treble gain, treble freq."],
            ["SET", "c4.amp.igain {input} {hex}", "Source leveling. AMP108: igain 01 ff annotated as −1 dB."],
            ["GET", "c4.amp.igain {input}", "Echoes input index and gain byte."],
            ["SET", "c4.amp.eq {zone} {five 6-digit groups}", "Five parametric bands. Q change used 00060e in the first group; defaults used 000605 000805 000b05 001105 001705."],
        ],
        mono_cols={0, 1},
    )
    doc.add_paragraph(
        "Signed gains are 8-bit two’s complement: 00 = 0 dB, 03 = +3, ff = −1. "
        "The five EQ groups are not fully decoded (frequency / gain / Q packing). "
        "Treat them as opaque unless you capture Composer while moving one slider at a time."
    )
    doc.add_paragraph("Composer controls with no recovered UDP string yet: Inputs Locked, Linear (dB) volume curve, Reset to Defaults as a single command.")

    doc.add_heading("5.4 Power save (model-dependent)", 2)
    doc.add_paragraph(
        "On AMP108, GET c4.amp.psave returned a mode byte; SET c4.amp.psave 00 woke the unit; "
        "02 disabled power save and 03 enabled it (status psave 00 02 / 00 03). "
        "On C4-16AMP3-B firmware 03.26.52, GET c4.amp.psave returns n01. Some Home Assistant scripts send two bytes (psave 00 00); the AMP108 capture used one byte (psave 00)."
    )

    doc.add_heading("5.5 Typical Director sequences (amplifier)", 2)
    doc.add_paragraph("Turn a zone on (Dining Room example from AMP108 log):")
    add_code(
        doc,
        "0g … c4.amp.digi 01\n0s … c4.amp.out 02 01\n0s … c4.amp.mute 02 00\n0s … c4.amp.chvol 02 a7",
    )
    doc.add_paragraph("Turn a zone off:")
    add_code(doc, "0s … c4.amp.mute {zone} 01\n0s … c4.amp.out {zone} 00")

    doc.add_heading("6. 16×16 audio switch — namespace c4.asw", 1)
    doc.add_paragraph(
        "Composer name: Control4 Audio Switch, type c4:v3_avswitch:avm-16s1-b "
        "(same family as C4-16ZAMSV3-B). Sixteen stereo RCA inputs and sixteen line-level RCA outputs. "
        "It does not drive speakers. Outputs feed the matrix amp and A/V receivers."
    )
    note(
        doc,
        "c4.amp.*, c4.switch.*, and c4.sw.* are the wrong family. The switch answers those with an empty 0r. The working prefix is c4.asw (audio switch).",
        "Important",
    )

    doc.add_heading("6.1 Routing, mute, and output gain", 2)
    add_table(
        doc,
        ["Get/Set", "Body", "Live result"],
        [
            ["GET", "c4.asw.ain", "16 hex bytes, input currently feeding each output. 00 = disconnected."],
            ["SET", "c4.asw.out {output} {input}", "000 on success. Example: out 10 01 routes input 1 to output 16; ain’s 16th byte becomes 01."],
            ["GET", "c4.asw.out", "e00 — GET is not valid without (or even with) this shape; use ain."],
            ["SET", "c4.asw.mute {output} 00|01", "000. GET amut updates that slot."],
            ["GET", "c4.asw.amut", "16 mute flags."],
            ["GET", "c4.asw.avol", "16 gain bytes. Factory/unity is 64 hex = 100 decimal."],
            ["SET", "c4.asw.vol {output} {hex}", "Output gain. 00 = silence/full cut, 64 = 100 / unity. chvol is n01 on this device."],
            ["GET", "c4.asw.abss / atrb / abal", "Bass, treble, balance (defaults 07, 07, 0c on this unit)."],
            ["GET", "c4.asw.psave", "n01"],
        ],
        mono_cols={0, 1},
    )

    doc.add_heading("6.2 Switch “volume” is line-level trim", 2)
    doc.add_paragraph(
        "Composer shows a Volume slider at 100, plus Bass, Treble, Balance, Source Leveling, "
        "Reset Settings, and Set Unity Gain. The switch is not a power amplifier. Volume is "
        "electronic output gain used to match levels between sources and destinations."
    )
    doc.add_paragraph("Confirmed behavior on unused output 16:")
    add_table(
        doc,
        ["Action", "Reply / ain or avol"],
        [
            ["SET c4.asw.vol 10 32 (50 decimal)", "000; avol 16th byte → 32"],
            ["SET c4.asw.vol 10 00", "000; avol 16th byte → 00"],
            ["SET c4.asw.vol 10 64 (100)", "000; avol 16th byte → 64 (unity)"],
            ["SET c4.asw.vol 10 78 (120 decimal)", "v01 but avol stayed 64 — values above 100 do not stick"],
        ],
        mono_cols={0, 1},
    )
    doc.add_paragraph(
        "So 100 (hex 64) is normal passthrough / unity. 0 is full attenuation. That is the "
        "right mental model for leveling different inputs and outputs. Speaker loudness still "
        "belongs on the 16AMP3 or the Sony receivers."
    )

    doc.add_heading("6.3 Worked routing example", 2)
    add_code(
        doc,
        "SET  0s6109 c4.asw.out 10 01\\r\\n\n"
        "RECV 0r6109 000\\r\\n\n"
        "GET  0g610a c4.asw.ain\\r\\n\n"
        "RECV 0r610a 000 c4.asw.ain 00 00 00 00 00 00 00 00 00 00 00 00 00 00 00 01\\r\\n\n"
        "SET  0s610b c4.asw.out 10 00\\r\\n\n"
        "RECV 0r610b 000\\r\\n",
    )

    doc.add_heading("7. Broken, rejected, and missing commands", 1)
    doc.add_paragraph(
        "This section is the working list of commands that fail, are unsafe, or exist in Composer "
        "but have no recovered UDP string yet. Empty 0r means the chassis parsed the packet but "
        "did not run a known handler. n01 means the name is in the protocol but not implemented "
        "on that firmware. e00 means the name is known but the arguments are wrong."
    )

    doc.add_heading("7.1 Do not use (broken or unsafe)", 2)
    add_table(
        doc,
        ["Command", "Product", "What happens", "Use instead"],
        [
            [
                "c4.amp.chvolmax",
                "Matrix amps",
                "Documented firmware bug: SET can snap the live zone volume to the cap instead of only storing a limit.",
                "Keep max volume in software; GET c4.amp.vlim is safe to read.",
            ],
            [
                "Volume = percent + 160",
                "Matrix amps",
                "Old community scripts. Reads about 5% high versus Director and hex 9b = 0%.",
                "percent + 155 (0% = 9b, 12% = a7).",
            ],
            [
                "c4.amp.mute {zone}   (no 00/01)",
                "Matrix amps",
                "AMP108 test: reply e00.",
                "c4.amp.mute {zone} 00|01",
            ],
            [
                "Zone 16 as decimal 16",
                "16x16 switch",
                "c4.asw.out 16 00 returned v01 but is not output 16. Output 16 is hex 10.",
                "Always two-digit hex: 10 = output 16.",
            ],
        ],
        mono_cols={0, 3},
    )

    doc.add_heading("7.2 Wrong namespace (switch vs amp)", 2)
    doc.add_paragraph(
        "Sending amplifier commands to the AVM-16S1-B produces an empty 0r (sequence only). "
        "The switch is alive; the name is simply not c4.amp."
    )
    add_table(
        doc,
        ["Tried on switch 192.168.68.80", "Reply"],
        [
            ["GET/SET c4.amp.ain, avol, out, mute, chvol, psave", "Empty 0r"],
            ["c4.switch.out / c4.switch.ain", "Empty 0r"],
            ["c4.sw.out / c4.sw.ain", "Empty 0r"],
            ["c4.avs.*, c4.avswitch.*, c4.avm.*, c4.ams.*, c4.matrix.*, c4.av.*", "Empty 0r"],
            ["c4.asw.chvol {zone} {hex}", "n01 — volume SET is c4.asw.vol, not chvol"],
            ["GET c4.asw.out", "e00 — list routes with GET c4.asw.ain"],
            ["GET c4.asw.psave", "n01"],
            ["c4.asw.ogain / gain / ovol / level", "n01"],
            ["SET c4.asw.vol 10 78 (above 100)", "v01 but avol unchanged — unity cap is 64 hex / 100"],
        ],
        mono_cols={0, 1},
    )

    doc.add_heading("7.3 Amplifier: rejected or model-specific", 2)
    add_table(
        doc,
        ["Command", "Where", "Reply"],
        [
            ["GET c4.amp.psave", "C4-16AMP3-B 03.26.52", "n01 (worked as GET/SET on AMP108 03.24.45)"],
            ["GET c4.sy.afwv", "AMP108", "n01"],
            ["GET c4.sy.mac / ver / ip", "AVM-16S1-B", "n01"],
        ],
        mono_cols={0, 1, 2},
    )

    doc.add_heading("7.4 Missing — Composer control with no recovered UDP yet", 2)
    doc.add_paragraph("Amplifier Advanced Properties (Composer) vs known packets:")
    add_table(
        doc,
        ["Composer control", "Status"],
        [
            ["Inputs Locked", "No UDP string captured."],
            ["Volume curve Linear (dB)", "No UDP string captured."],
            ["Reset to Defaults (one button)", "Unknown as a single command; Director may replay many SET lines."],
            ["Maximum Volume SET that is safe", "GET vlim works. Hardware SET chvolmax is unsafe (see 7.1). Unsolicited chlim exists after volume ramps."],
            ["Parametric EQ 5 bands (Hz, gain, Q)", "SET c4.amp.eq exists (six-digit groups) but frequency/gain/Q packing is not fully decoded. Q=10 used 00060e in band 1."],
            ["EQ Reset to Defaults / Basic", "Not isolated; likely a batch of eq + bass/treble SET."],
            ["Digital input select (beyond GET digi)", "GET c4.amp.digi is used before route; SET form not confirmed on 16AMP3."],
        ],
    )
    doc.add_paragraph("Audio switch Advanced Properties vs known packets:")
    add_table(
        doc,
        ["Composer control", "Status"],
        [
            ["Output Volume 0-100", "Confirmed: GET avol, SET vol. 100 = unity."],
            ["Mute", "Confirmed: SET mute, GET amut."],
            ["Bass / Treble / Balance sliders", "GET abss, atrb, abal work. SET names (bassgain / trebgain / bal) not live-tested on the switch."],
            ["Source Leveling (input, dB, Set)", "Likely c4.asw.igain by analogy with c4.amp.igain; not confirmed on AVM-16S1-B."],
            ["Set Unity Gain", "Unknown as one command; SET vol {n} 64 on each output would match 100."],
            ["Reset Settings", "Unknown as one command."],
            ["DHCP / static IP Apply Settings", "Not in the c4.asw audio namespace; network is Composer/SDDP, not these UDP strings."],
        ],
    )

    doc.add_heading("7.5 Captured but not fully decoded", 2)
    add_table(
        doc,
        ["Command", "What we know", "What is missing"],
        [
            ["c4.amp.eq {zone} five groups", "Accepted 000. Defaults and Q change captured on AMP108.", "Which hex fields are Hz, gain, Q for bands 80/125/250/1k/4k."],
            ["c4.amp.bassfreq / trebfreq", "Director sent 08 and 18; tone status 00 08 00 18.", "Exact Hz mapping to Composer 125 Hz / 5 kHz."],
            ["c4.amp.igain", "ff annotated -1 dB on AMP108.", "Full dB table; 16AMP3 not live-tested."],
            ["c4.asw.abss = 07, abal = 0c", "Idle defaults on all 16 outputs.", "How 07/0c map to Composer 0 dB and center C."],
            ["c4.asw.outvol", "SET returned v01; avol not verified against it.", "Whether it is an alias of vol or a no-op."],
        ],
        mono_cols={0},
    )

    doc.add_heading("8. Typical physical graph (example site)", 1)
    add_table(
        doc,
        ["Switch jack", "Connected to"],
        [
            ["Input 1", "EA-5 audio out 1"],
            ["Input 3", "Retro Hi-Fi"],
            ["Input 4", "WiiM Pro"],
            ["Output 1", "Amplifier audio input 1"],
            ["Output 2", "Basement receiver SA-CD/CD"],
            ["Output 3", "Living Room receiver SA-CD/CD"],
            ["Outputs 4–16", "Unused (safe for SET tests)"],
        ],
    )
    add_table(
        doc,
        ["Amp zone", "Room (Composer Room Selection)"],
        [
            ["1", "Kitchen"],
            ["2", "Dining Room"],
            ["3", "Patio"],
            ["4", "Master Bedroom"],
            ["5", "Master Bathroom"],
            ["6", "Bar"],
            ["7", "Garage"],
            ["8", "Unused in Composer"],
        ],
    )
    note(
        doc,
        "Composer bindings are logical. Live GET c4.amp.ain can still be all zeros if no zone is currently routed, even though rooms are mapped.",
    )

    doc.add_heading("9. How to test without Composer", 1)
    doc.add_paragraph("Probe (GET-only identify, then optional SET):")
    add_code(
        doc,
        "python tools/c4_probe.py --host 192.168.68.93 --identify\n"
        "python tools/c4_probe.py --host 192.168.68.80 --get \"c4.asw.ain\"",
    )
    doc.add_paragraph(
        "Relay / intermediary: bind UDP 8750 on a PC, point Composer’s device IP at that PC, "
        "forward to the real chassis, log both directions. Same idea as Scripts/UDP_intermediary.py "
        "and tools/c4_relay.py."
    )

    doc.add_heading("10. Quick command cheat sheet", 1)
    add_table(
        doc,
        ["Task", "Amplifier", "Audio switch"],
        [
            ["Identify", "c4.sy.fwv / c4.sy.info", "same"],
            ["List routes", "GET c4.amp.ain", "GET c4.asw.ain"],
            ["Route", "SET c4.amp.out zz ii", "SET c4.asw.out zz ii"],
            ["Disconnect", "SET c4.amp.out zz 00", "SET c4.asw.out zz 00"],
            ["Mute", "SET c4.amp.mute zz 01", "SET c4.asw.mute zz 01"],
            ["Volume / gain", "SET c4.amp.chvol zz {pct+155}", "SET c4.asw.vol zz {0–100 as hex}"],
            ["Wrong family looks like", "—", "Empty 0r if you send c4.amp.*"],
        ],
        mono_cols={1, 2},
    )

    doc.add_heading("11. Sources", 1)
    bullets = [
        "AMP108 Director ↔ amp UDP log (C4-AMP108-1B Scripts, July 2024), including annotated igain / eq / psave.",
        "Live GET/SET on C4-16AMP3-B 192.168.68.93 and AVM-16S1-B 192.168.68.80 (August 2026).",
        "Composer Pro Advanced Properties (amp EQ, switch volume/unity gain).",
        "Community Home Assistant UDP amp integrations (framing, chvol +155, chvolmax caveat).",
    ]
    for item in bullets:
        doc.add_paragraph(item, style="List Bullet")

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.text = "Control4 Ethernet Audio — unofficial UDP reference  ·  Not a Control4 publication"
    fp.runs[0].font.size = Pt(8)
    fp.runs[0].font.color.rgb = RGBColor(127, 127, 127)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
